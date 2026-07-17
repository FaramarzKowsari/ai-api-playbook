from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


@dataclass(frozen=True)
class PageSpec:
    number: int
    title: str
    subtitle: str
    visual: str
    points: tuple[str, str, str, str]
    code: str
    input_text: str
    output_text: str
    commercial: str
    repo_path: str


def p(number: int, title: str, subtitle: str, visual: str, points: tuple[str, str, str, str], code: str, input_text: str, output_text: str, commercial: str, repo_path: str) -> PageSpec:
    return PageSpec(number, title, subtitle, visual, points, code, input_text, output_text, commercial, repo_path)


PAGES = [
    p(1,"THE AI API ECOSYSTEM","One request can activate text, tools, knowledge, sight, voice and video.","A friendly cartoon engineer at the center of a six-orbit capability solar system",("Text APIs generate, reason and code","Tool APIs connect models to actions","Retrieval APIs ground answers in evidence","Media APIs understand and create multiple modalities"),"request → model + tools + data → validated result","Customer intent + authorized context","Useful result + evidence + usage metrics","Value appears when a capability solves a recurring customer problem.","docs/book-map.md"),
    p(2,"ANATOMY OF AN AI API CALL","Every reliable call has identity, instructions, input, limits and a destination.","Exploded-view cartoon of a transparent API request capsule",("Endpoint chooses the service","Authentication proves permitted access","Payload carries messages, tools and constraints","Timeout and request ID make failures manageable"),'POST /v1/...\nAuthorization: Bearer $KEY\nContent-Type: application/json',"Prompt, model policy, schema, metadata","Content, tool calls, usage, status, request ID","A reusable request wrapper reduces engineering time across products.","src/ai_api_playbook/models.py"),
    p(3,"ENDPOINTS, METHODS & HEADERS","The URL says where; the method says what; headers say how.","Color-coded subway map where endpoint, method, headers and body are stations",("GET reads; POST creates or computes","Headers carry auth, content type and trace data","Path parameters identify resources","Query parameters tune filtering and pagination"),'POST https://api.example.com/v1/responses',"Headers + JSON body","HTTP status + headers + JSON body","A clean API client becomes reusable infrastructure for many paid services.","docs/architecture.md"),
    p(4,"API KEY SECURITY","A key is not a password you type everywhere—it is a capability you must contain.","Cartoon vault with server inside and browser outside",("Store secrets in a managed secret store","Never commit .env or expose privileged keys to clients","Rotate leaked credentials immediately","Use least privilege, spend limits and separate environments"),'key = os.environ["PROVIDER_API_KEY"]',"Server-side secret reference","Authenticated request without revealing the secret","Security is part of product value, not an afterthought.","SECURITY.md"),
    p(5,"THE REQUEST PAYLOAD","Good payloads remove ambiguity before the model begins.","Bento-box JSON payload with labeled compartments",("Separate instructions from user data","Set explicit output limits","Use schemas for machine-consumed results","Attach metadata for tracing, not sensitive content"),'{"model":"...","input":"...","max_output_tokens":800}',"Instructions + content + constraints","A deterministic request contract","Fewer retries and manual corrections improve gross margin.","src/ai_api_playbook/models.py"),
    p(6,"READING THE RESPONSE","Do not confuse generated text with a completed, validated business result.","Airport arrivals board showing status, content, usage and request ID",("Check completion status and finish reason","Parse tool calls separately from text","Record token usage and latency","Validate structure before saving or acting"),'if response.status != "completed": raise RuntimeError()',"Provider response object","Validated domain object + safe telemetry","Reliable response handling reduces support cost and failed jobs.","src/ai_api_playbook/client.py"),
    p(7,"REST: THE DEFAULT TRANSPORT","Simple request/response remains ideal for short independent jobs.","Restaurant order ticket metaphor with client and server",("REST is easy to test, cache and observe","Use POST for generation requests","Set connect and read timeouts","Retry only safe or idempotent operations"),'r = httpx.post(url, json=payload, timeout=30)',"One complete request","One complete response","REST is the fastest path to a sellable backend MVP.","examples/01_text/openai_responses.py"),
    p(8,"STREAMING WITH SSE","Show useful output before the model finishes the full answer.","River of token cards flowing from cloud to screen",("Streaming reduces perceived latency","Process typed events, not raw text fragments","Accumulate a final canonical result","Handle disconnects and partial output explicitly"),'for event in stream:\n    render(event.delta)',"Long generation request","Incremental events + final completion","Faster perceived response improves conversion in interactive products.","docs/architecture.md"),
    p(9,"WEBSOCKET & WEBRTC","Realtime systems keep a living connection instead of repeating isolated calls.","Two-way neon bridge carrying audio waves and tool events",("WebSocket supports bidirectional event streams","WebRTC is optimized for browser media","Use ephemeral client credentials","Design interruption, reconnection and session expiry"),'wss://provider.example/v1/realtime',"Continuous audio/events","Continuous audio, transcripts and tool calls","Realtime voice enables premium, usage-based services.","examples/05_media/deepgram_stt.py"),
    p(10,"ASYNCHRONOUS JOBS","Media generation is a job lifecycle, not a long frozen HTTP request.","Cartoon factory conveyor: queued, processing, succeeded, failed",("Create a job and store its ID","Poll politely or receive a webhook","Expose progress without promising exact timing","Make retries idempotent to prevent duplicate charges"),'job = create_task()\nstatus = get_task(job.id)',"Prompt + media references","Job ID → status → output URL","Async queues make expensive media products scalable.","examples/05_media/runway_video.py"),
    p(11,"WEBHOOKS","Let the provider notify your server when a long task changes state.","Doorbell metaphor with signed envelope arriving at a server",("Verify signature and timestamp","Acknowledge quickly; process in a queue","Deduplicate repeated deliveries","Never trust payload URLs or IDs without validation"),'verify(signature, timestamp, raw_body)',"Signed event delivery","200 acknowledgement + queued work","Webhooks reduce polling cost and improve operational reliability.","SECURITY.md"),
    p(12,"BATCH PROCESSING","Trade immediate answers for lower cost and higher throughput.","Night-shift robot processing a stack of JSONL cards",("Batch independent, non-urgent work","Use stable custom IDs for reconciliation","Validate every line before upload","Plan for partial failures and completion windows"),'custom_id, method, url, body',"JSONL request file","JSONL results matched by custom ID","Batch is ideal for catalogs, archives and recurring reports.","docs/architecture.md"),
    p(13,"LANGUAGE MODELS AS ENGINES","The model predicts output; your system defines purpose, evidence and control.","Cartoon engine cutaway labeled reasoning, language and code",("Models transform context into probabilistic output","The same model can classify, extract, draft and explain","Quality depends on task definition and evidence","Business reliability requires validation outside the model"),'response = client.generate(request)',"Messages + constraints","Text or structured result","Sell the solved workflow—not access to a generic model.","src/ai_api_playbook/client.py"),
    p(14,"MODEL SELECTION","Choose with measurements, not reputation.","Four-lane racing dashboard for quality, latency, cost and context",("Start with task-specific evaluation cases","Measure quality per successful job","Include latency and failure rate","Route premium and routine work differently"),'choose_model(candidates, strategy="balanced")',"Candidate metrics + product policy","Selected provider/model","Routing protects margin while preserving quality.","src/ai_api_playbook/routing.py"),
    p(15,"TOKENS & COST","Characters become tokens; tokens become latency and variable cost.","Token coins moving through input and output meters",("Input and output may have different prices","Tools and retrieved context also consume tokens","Estimate before calling and measure after","Limit unnecessary context and verbose output"),'estimated = input_tokens*in_rate + output_tokens*out_rate',"Prompt, context, output budget","Usage record + estimated cost","Cost visibility enables profitable pricing.","src/ai_api_playbook/economics.py"),
    p(16,"THE CONTEXT WINDOW","More context is capacity, not permission to send everything.","A suitcase packing metaphor with evidence and clutter",("Keep only task-relevant evidence","Place durable policy separately from user data","Summarize or compact long histories","Never treat context length as retrieval quality"),'context = select_relevant(chunks, budget)',"Documents + conversation + token budget","Focused evidence set","Lean context improves speed, cost and accuracy.","src/ai_api_playbook/rag.py"),
    p(17,"PROMPT ENGINEERING","A strong prompt is a specification with acceptance criteria.","Architect blueprint with goal, context, constraints and tests",("State the goal and audience","Provide authoritative context","Define constraints and forbidden behavior","Specify output shape and success checks"),'GOAL + CONTEXT + CONSTRAINTS + FORMAT + CHECKS',"Task specification","Testable response","Reusable prompts become product assets when versioned and evaluated.","docs/architecture.md"),
    p(18,"INSTRUCTION HIERARCHY","Policies, developer intent, user input and retrieved data are not equal.","Layered security tower with four colored floors",("Keep system policy stable and minimal","Developer instructions define workflow","User input supplies the task","Retrieved content is untrusted evidence, not instruction"),'policy > workflow > user task > external data',"Mixed instruction sources","Resolved behavior with boundaries","Clear hierarchy reduces costly and dangerous agent mistakes.","SECURITY.md"),
    p(19,"STRUCTURED OUTPUTS","If software will consume the answer, ask for a contract—not prose.","JSON crystal growing inside a schema frame",("Define types, required fields and allowed values","Disallow unexpected properties when appropriate","Validate again in application code","Handle refusal and invalid output paths"),'schema={"type":"object","required":["name"]}',"Unstructured content","Schema-valid JSON object","Structured extraction powers automation and data products.","examples/01_text/openai_responses.py"),
    p(20,"JSON SCHEMA DESIGN","The schema is a user interface for the model and a safety boundary for code.","Cartoon customs checkpoint inspecting JSON fields",("Prefer small explicit objects","Use enums for bounded choices","Add descriptions that clarify meaning","Represent uncertainty with confidence and evidence fields"),'{"status":{"enum":["pass","review","fail"]}}',"Raw model candidate","Validated typed record","A good schema reduces downstream exceptions and review time.","src/ai_api_playbook/models.py"),
    p(21,"REASONING WORKLOADS","Give difficult decisions enough budget, tools and evidence—then test outcomes.","Mountain route with checkpoints and a final verifier",("Separate reasoning quality from verbosity","Use tools for current or private facts","Provide explicit decision criteria","Evaluate answers, not hidden thought processes"),'result = solve(problem, evidence, effort="high")',"Complex task + evidence","Decision + concise rationale + citations","Premium reasoning is valuable when tied to high-value decisions.","src/ai_api_playbook/evaluation.py"),
    p(22,"CODE GENERATION","Generated code becomes trustworthy only after tests, review and constrained execution.","Robot programmer passing gates: lint, tests, security, review",("Provide repository conventions and interfaces","Generate the smallest scoped change","Run static checks and tests","Review dependencies, permissions and failure modes"),'change → compile → test → review → merge',"Issue + code context","Patch + verification evidence","Code assistants can become productized modernization services.","CONTRIBUTING.md"),
    p(23,"MULTIMODAL INPUT","One request can combine text, images, audio, video and documents.","Five colored streams merging into one reasoning lens",("Label each asset and its purpose","Preserve source and page/time provenance","Control media resolution and size","Ask cross-modal questions that require the evidence"),'input=[text, image, audio, document]',"Mixed authorized media","Unified analysis with source references","Multimodal workflows unlock richer vertical products.","docs/provider-matrix.md"),
    p(24,"MODEL ROUTING & FALLBACK","Fallback should preserve the product contract, not merely return any answer.","Air-traffic control tower routing jobs among model runways",("Route by measured task performance","Keep schemas consistent across providers","Use circuit breakers for failing endpoints","Do not silently downgrade critical workflows"),'primary → retry → compatible fallback → human queue',"Task policy + live metrics","Best eligible route + audit record","Multi-provider resilience can become an enterprise differentiator.","src/ai_api_playbook/routing.py"),
    p(25,"FUNCTION CALLING","The model proposes arguments; your application decides whether and how to execute.","Cartoon model handing a sealed tool request to a strict gatekeeper",("Define narrow tools with clear descriptions","Validate every argument against a schema","Authorize by user, tenant and purpose","Return typed tool results to the model"),'tool_call = {"name":"quote_price","arguments":{...}}',"User intent + tool catalog","Structured tool proposal","Tool-enabled assistants solve workflows instead of only answering questions.","examples/02_tools/function_calling.py"),
    p(26,"DESIGNING TOOL SCHEMAS","Small tools are easier to authorize, observe and repair.","Workbench with labeled single-purpose tools",("One tool should do one bounded job","Use enums and numeric limits","Describe side effects explicitly","Separate read tools from write tools"),'name + description + input_schema',"Permitted operation","Validated call with predictable effect","A reusable tool catalog accelerates multiple agent products.","examples/02_tools/mcp_tool_contract.json"),
    p(27,"THE AGENT LOOP","Observe, decide, act, verify—within a budget and stop condition.","Circular four-stage board game with a budget meter",("Start from an explicit goal and state","Allow only necessary tools","Check results after every action","Stop on success, limit, uncertainty or human review"),'while not done and steps < limit: observe(); act(); verify()',"Goal + state + tool permissions","Completed result or controlled escalation","Agents create value by completing bounded processes.","docs/architecture.md"),
    p(28,"CLIENT TOOLS VS SERVER TOOLS","Execution location changes trust, latency, cost and responsibility.","Split-screen kitchen: client-side chef and provider-side chef",("Client tools run in your controlled application","Server tools run on provider infrastructure","Record where data and code travel","Apply least privilege in both locations"),'tool.execution = "client" | "server"',"Tool request","Tool result + execution metadata","Choosing execution boundaries is an enterprise architecture service.","docs/provider-matrix.md"),
    p(29,"WEB SEARCH TOOLS","Search retrieves candidates; grounded synthesis must still inspect and cite.","Cartoon detective following search clues to source cards",("Use search for unstable facts","Prefer authoritative primary sources","Track URL, date and claim support","Separate source evidence from model inference"),'search(query) → inspect(results) → synthesize(citations)',"Current question","Answer with directly supporting links","Cited research briefs are sellable recurring products.","docs/architecture.md"),
    p(30,"FILE SEARCH","Turn approved files into a governed evidence layer.","Library shelves becoming searchable vector cards",("Define ingestion ownership and retention","Chunk with page and section metadata","Filter by access scope before ranking","Return citations that a reader can verify"),'retrieve(query, tenant_id, permissions)',"Question + authorized corpus","Ranked passages + source IDs","Private knowledge assistants have strong recurring value.","examples/03_rag/local_rag.py"),
    p(31,"CODE EXECUTION TOOLS","Use sandboxes for calculation and transformation—not unrestricted trust.","Glass sandbox containing a tiny Python robot",("Constrain filesystem, network, time and memory","Pin dependencies and inspect inputs","Capture stdout, files and exit status","Never mix secrets with untrusted code"),'run(code, timeout=30, network=False)',"Code + safe files","Logs, exit code and artifacts","Controlled analysis environments enable premium data services.","SECURITY.md"),
    p(32,"COMPUTER USE","Pixels are flexible but fragile; APIs remain preferable when available.","Robot hand using mouse beside a clean API plug",("Use APIs for stable structured operations","Computer use needs visual verification","Confirm destructive or external actions","Design recovery for popups, layout changes and stale state"),'observe → locate → act → verify',"Screenshot + task","Verified UI state change","UI automation can bridge legacy systems with careful supervision.","SECURITY.md"),
    p(33,"MODEL CONTEXT PROTOCOL","MCP standardizes how models discover and call external tools and resources.","Universal adapter connecting many services to one agent",("Expose clear tool contracts and resources","Authenticate users and preserve authorization","Keep server permissions narrow","Treat remote results as untrusted input"),'agent ↔ MCP client ↔ MCP server ↔ service',"Tool discovery or call","Structured result with provenance","A well-designed MCP server can distribute a specialized business capability.","examples/02_tools/mcp_tool_contract.json"),
    p(34,"MULTI-AGENT SYSTEMS","Add agents only when specialization or parallel work beats coordination cost.","Small expert cartoon team around a shared task board",("Give each agent a distinct role and boundary","Share explicit state, not vague conversation","Use one owner for final synthesis","Measure handoff errors and token overhead"),'planner → specialists → reviewer → final owner',"Complex decomposable task","Integrated result + trace","Specialist workflows can support high-value research and operations.","docs/architecture.md"),
    p(35,"EMBEDDINGS","An embedding maps meaning into a vector that software can compare.","Words turning into points in a colorful 3D galaxy",("Similar meanings tend to have nearby vectors","Use the same compatible model for corpus and query","Normalize and version embedding pipelines","Embeddings enable search, clustering and recommendations"),'vector = embed("semantic meaning")',"Text or supported content","Numeric vector","Semantic infrastructure powers reusable knowledge products.","src/ai_api_playbook/rag.py"),
    p(36,"VECTOR SPACE INTUITION","Distance becomes a practical proxy for semantic relatedness.","Treasure map with clusters, arrows and nearest neighbors",("Dimensions are learned features, not human labels","Cosine compares direction","Distance alone does not prove relevance","Evaluate retrieval on real user questions"),'cosine = dot(a,b)/(norm(a)*norm(b))',"Query vector + document vectors","Similarity scores","Good retrieval increases answer quality without premium generation every time.","src/ai_api_playbook/rag.py"),
    p(37,"COSINE SIMILARITY","Compare direction when magnitude should not dominate meaning.","Two compass arrows with angle highlighted",("A score near one means similar direction","Zero vectors require explicit handling","Thresholds are task-specific","Ranking quality matters more than an arbitrary universal cutoff"),'cosine_similarity([1,0],[1,0]) == 1.0',"Two equal-length vectors","Similarity score","Transparent similarity math builds technical credibility.","tests/test_core.py"),
    p(38,"CHUNKING","The chunk is the unit of evidence your retriever can find and your reader can verify.","Document cut into puzzle pieces with small overlaps",("Split along semantic boundaries where possible","Use overlap to preserve cross-boundary meaning","Keep page and heading metadata","Test chunk size against retrieval cases"),'chunks = chunk_text(text, target_chars=800, overlap_chars=120)',"Long document","Traceable evidence chunks","Better chunking reduces costly re-generation and manual search.","src/ai_api_playbook/rag.py"),
    p(39,"METADATA FILTERS","Filter permissions and scope before similarity ranking.","Security turnstile before a vector-search racetrack",("Store tenant, language, date, type and access metadata","Apply authorization before retrieval","Use filters to reduce irrelevant candidates","Do not trust model-generated metadata without validation"),'filter={"tenant_id":{"$eq": tenant}}',"Query + access scope","Eligible candidate set","Governed retrieval is essential for enterprise subscriptions.","examples/07_search/pinecone_vector_store.py"),
    p(40,"VECTOR DATABASES","A vector database stores embeddings plus metadata for fast similarity search.","Cartoon warehouse of vector boxes with a search crane",("Choose managed or self-hosted by scale and governance","Plan index dimension and metric","Support namespaces or tenant isolation","Track deletion and re-embedding workflows"),'index.upsert(vectors); index.query(vector, top_k=5)',"Vectors + IDs + metadata","Ranked matching records","Managed retrieval can anchor recurring knowledge services.","examples/07_search/pinecone_vector_store.py"),
    p(41,"THE RAG PIPELINE","Retrieve evidence first; generate only from the evidence that survives ranking.","Six-stage water-filtration pipeline labeled ingest to cite",("Ingest and normalize authorized sources","Chunk, embed and index with provenance","Retrieve, filter and rerank","Generate with evidence and return citations"),'ingest → chunk → embed → retrieve → rerank → answer',"Question + governed corpus","Grounded answer + sources","RAG converts private knowledge into a controlled service.","examples/03_rag/local_rag.py"),
    p(42,"HYBRID SEARCH","Combine semantic understanding with exact lexical matching.","Two rivers—keyword and vector—merging through reciprocal-rank fusion",("Keywords catch names, codes and exact phrases","Vectors catch paraphrases and concepts","Fuse rankings rather than raw incompatible scores","Evaluate by query category"),'RRF = Σ 1/(k + rank)',"Keyword ranking + vector ranking","Fused candidates","Hybrid search improves reliability in technical and catalog domains.","src/ai_api_playbook/rag.py"),
    p(43,"RERANKING","Spend a stronger relevance model only on the small candidate set that matters.","Talent-show judge reordering ten candidate cards",("Retrieve broadly and cheaply","Rerank top candidates with query-aware scoring","Pass only the best evidence to generation","Measure precision and end-to-end answer quality"),'rerank(query, candidates[:50])[:5]',"Query + candidate passages","Improved final evidence order","Reranking can lift quality without paying for a larger generator.","examples/07_search/cohere_rerank.py"),
    p(44,"RAG EVALUATION","A fluent answer can still be unsupported, incomplete or badly cited.","Four-gauge quality dashboard",("Measure retrieval recall separately","Score groundedness against provided evidence","Check citation correctness and completeness","Track latency, cost and abstention behavior"),'quality = .35*grounded + .30*complete + ...',"Test questions + expected evidence","Quality, latency and cost report","Evaluation is the difference between a demo and a dependable product.","src/ai_api_playbook/evaluation.py"),
    p(45,"COMPUTER VISION APIS","Vision models translate pixels into descriptions, fields, decisions and actions.","Cartoon camera lens splitting into caption, boxes, table and alert cards",("Choose resolution for the task","Ask precise visual questions","Preserve original asset and provenance","Do not infer hidden traits from appearance"),'analyze(image, task="inventory_count")',"Authorized image + task","Description, objects or structured fields","Vertical vision products can automate expensive inspection workflows.","docs/provider-matrix.md"),
    p(46,"IMAGE UNDERSTANDING","Point the model toward observable evidence, not speculation.","Magnifying glass over a product shelf with evidence labels",("Separate observation from inference","Request coordinates or evidence descriptions","Use multiple views for physical tasks","Human-review high-impact conclusions"),'{"observations":[],"uncertainties":[],"confidence":0.0}',"Image + question + schema","Auditable visual observations","Evidence-first vision increases customer trust.","docs/architecture.md"),
    p(47,"OCR & DOCUMENT AI","OCR should recover structure, not just a stream of characters.","Paper invoice unfolding into headings, tables and bounding boxes",("Preserve pages, reading order and hierarchy","Extract tables and coordinates when needed","Use confidence to create a review queue","Keep source-page provenance for every field"),'POST /v1/ocr → pages[].markdown + blocks[]',"PDF or scan","Markdown, tables, blocks and coordinates","Document extraction saves measurable manual re-keying time.","examples/04_documents/mistral_ocr.py"),
    p(48,"PDF UNDERSTANDING","A PDF can contain text, scans, tables, charts and misleading reading order.","X-ray view of a PDF showing multiple hidden layers",("Detect native text versus scanned pages","Render pages when layout carries meaning","Chunk by heading and page","Validate tables, equations and footnotes separately"),'pdf → inspect → OCR/render → structure → verify',"Authorized PDF","Structured document with page references","Reliable PDF pipelines serve publishing, accounting and research.","projects/publishing_engine/pipeline.py"),
    p(49,"STRUCTURED DOCUMENT EXTRACTION","Every extracted value needs a type, evidence location and confidence path.","Customs form with green validated fields and amber review fields",("Define a strict domain schema","Attach page and bounding-box evidence","Normalize dates, currencies and identifiers","Send low-confidence or conflicting fields to review"),'{"value":"...","page":3,"confidence":0.91}',"Document pages","Validated record + exception queue","Charge for successful verified documents, not raw OCR pages alone.","projects/blueprints.py"),
    p(50,"IMAGE GENERATION APIS","Creative output becomes useful when the brief controls purpose, format and review.","Prompt storyboard transforming into three polished asset variants",("Specify subject, composition, medium and constraints","Use references only with proper rights","Generate variants for selection","Inspect text, anatomy, brands and factual claims"),'generate(prompt, size="A4", variants=3)',"Creative brief + authorized references","Candidate images + metadata","Productized creative services can sell outcome packages.","examples/05_media/openai_image_generation.py"),
    p(51,"IMAGE EDITING","Editing requires both a source image and an explicit change boundary.","Before/after image with a glowing masked region",("State what must remain unchanged","Use masks or regional instructions","Preserve aspect ratio and identity only with consent","Compare output against the source at full resolution"),'edit(image, mask, instruction)',"Source + mask + edit brief","Edited candidate","Controlled editing supports catalogs, marketing and restoration workflows.","examples/05_media/openai_image_generation.py"),
    p(52,"VISUAL QUALITY ASSURANCE","Generated does not mean publishable.","Cartoon quality inspector with zoom lens and checklist",("Check spelling and typography at 100% zoom","Verify logos, products and factual labels","Inspect crops, hands, faces and repeated elements","Record seed, prompt, model and human approval"),'generate → render → inspect → revise → approve',"Candidate visual","Approved asset + provenance","QA protects brand reputation and reduces refunds.","SECURITY.md"),
    p(53,"SPEECH TO TEXT","Transcription turns audio into searchable, time-aligned data.","Audio waveform becoming timestamped dialogue cards",("Choose prerecorded or streaming mode","Request timestamps and diarization when needed","Handle accents, noise and domain vocabulary","Keep confidence and original audio for review"),'POST /v1/listen?diarize=true',"Audio stream or file","Transcript + speakers + timestamps","Transcription supports meetings, media and support analytics.","examples/05_media/deepgram_stt.py"),
    p(54,"TEXT TO SPEECH","A voice API converts approved text into controlled audio delivery.","Script cards flowing into a friendly studio microphone",("Choose language, voice and speaking style","Stream when text arrives incrementally","Normalize numbers, dates and abbreviations","Disclose synthetic speech where appropriate"),'POST /v1/text-to-speech/{voice_id}',"Text + voice settings","Audio bytes or stream","TTS enables audiobooks, learning and accessible content.","examples/05_media/elevenlabs_tts.py"),
    p(55,"REALTIME VOICE","A natural voice product is a latency system before it is a language demo.","Circular waveform loop: listen, detect turn, think, speak",("Measure time to first audio","Support interruption and barge-in","Keep session state bounded","Use tools without blocking the conversation"),'audio in ↔ realtime session ↔ audio out',"Live microphone audio","Transcript, tool events and speech","Realtime voice supports premium assistants and reception services.","docs/provider-matrix.md"),
    p(56,"VOICE AGENT ARCHITECTURE","Listening, reasoning, tools and speaking must share one controlled session.","Four-module cartoon call center around a shared state hub",("Authenticate the caller and session","Separate STT, reasoning, tools and TTS metrics","Confirm names, dates and irreversible actions","Escalate uncertainty to a person"),'STT → agent → tool → TTS',"Caller speech","Answer, booking result and call summary","Voice agents monetize through setup, subscription and minutes.","projects/blueprints.py"),
    p(57,"TURN-TAKING & INTERRUPTION","The best answer fails if the agent speaks over the customer.","Conversation traffic light with listen, think and speak states",("Detect end of turn without long silence","Allow user interruption","Cancel unheard audio and stale tool plans","Log what the caller actually heard"),'LISTENING → THINKING → SPEAKING ↔ INTERRUPTED',"Audio timing events","Natural turn transitions","Better turn-taking increases completion rate and customer satisfaction.","docs/architecture.md"),
    p(58,"TRANSLATION & DUBBING","Localization preserves meaning, timing, tone and cultural fit—not just words.","Film strip crossing language bridges with timed audio bubbles",("Transcribe with timestamps first","Translate for audience and context","Generate or cast voice with consent","Review names, numbers, claims and lip timing"),'audio → transcript → translation → voice → mix → review',"Authorized media + target locale","Timed localized media","Localization is a productized service with clear per-minute pricing.","projects/blueprints.py"),
    p(59,"TELEPHONY INTEGRATION","A phone call adds networks, codecs, identity and operational risk.","Phone network diagram connecting caller, carrier, websocket and agent",("Validate provider webhooks and call identity","Convert audio formats deliberately","Handle transfer, voicemail and disconnects","Record only with lawful notice and consent"),'PSTN → telephony provider → media stream → voice agent',"Inbound call","Resolution, transfer or safe fallback","Missed-call recovery has measurable commercial value.","projects/blueprints.py"),
    p(60,"VOICE SAFETY & CONSENT","A realistic voice is a capability with identity and fraud risk.","Consent shield around a voice waveform and identity card",("Obtain explicit permission for cloning","Prevent impersonation and deceptive use","Protect voice samples as sensitive data","Provide disclosure, revocation and audit trails"),'consent.valid && purpose.allowed',"Voice sample + authorized purpose","Bounded synthetic voice use","Trustworthy voice services can compete on governance.","SECURITY.md"),
    p(61,"VIDEO GENERATION APIS","Video is an expensive sequence of consistent frames, not a single image.","Storyboard conveyor turning text cards into moving frames",("Define shot, motion, duration and aspect ratio","Use reference assets with rights","Budget per second and per retry","Review continuity, text, identity and claims"),'create_video(prompt, duration=5, ratio="16:9")',"Prompt + optional image/video references","Asynchronous video task","Short-form product packages can be sold per approved clip.","examples/05_media/runway_video.py"),
    p(62,"IMAGE TO VIDEO","A strong first frame anchors identity; motion instructions control the transformation.","Still illustration opening into a cinematic motion tunnel",("Use a clean high-resolution source","Describe camera and subject motion separately","Avoid contradictory motion","Compare first and last frames for drift"),'image + motion_prompt → video job',"Authorized image + motion brief","Short generated clip","Image-to-video upgrades existing catalogs and campaigns.","examples/05_media/runway_video.py"),
    p(63,"VIDEO EDITING WITH AI","Editing should name the exact transformation and protected content.","Timeline with highlighted editable segments and locked regions",("Specify timestamps and intended change","Protect identity, product and brand details","Use keyframes when supported","Review every frame near transitions"),'edit(video, instruction, keyframes)',"Video + edit boundary","Edited video candidate","AI-assisted editing can reduce turnaround for media teams.","docs/provider-matrix.md"),
    p(64,"AVATARS & DIGITAL PRESENTERS","A presenter is a production layer—not permission to imitate a person.","Friendly digital presenter beside a consent contract and script",("Use owned or licensed likenesses","Bind avatar to approved scripts and knowledge","Disclose synthetic representation","Monitor identity drift and prohibited claims"),'avatar + approved_audio + disclosure',"Consented identity + script","Presenter video","Avatars support scalable training and multilingual explainers.","SECURITY.md"),
    p(65,"MUSIC & SOUND EFFECTS","Sound generation needs mood, timing, format and rights-aware review.","Colorful soundboard with mood, duration and loop controls",("Define purpose, duration, energy and instrumentation","Avoid imitation requests involving living artists","Generate stems or loops when workflow needs them","Check loudness, clipping and usage rights"),'generate_sound("soft interface confirmation", seconds=2)',"Audio brief","Sound effect or music candidate","Licensed sound packages can complement video and learning products.","docs/provider-matrix.md"),
    p(66,"THE MEDIA JOB PIPELINE","Upload, create, monitor, validate, deliver and expire.","Six-stage media production factory with status lights",("Use secure temporary asset URLs","Store provider task IDs and idempotency keys","Queue polling and webhook work","Expire outputs according to retention policy"),'upload → submit → monitor → validate → deliver → delete',"Media assets + brief","Approved output + provenance + cost","Operational discipline protects margin in media products.","examples/05_media/runway_video.py"),
    p(67,"SECRETS & IDENTITY","Production authentication should be short-lived, scoped and observable.","Identity badge passing through layers of a zero-trust castle",("Prefer workload identity where supported","Separate development, staging and production","Rotate and revoke automatically","Alert on unusual spend, region or request volume"),'credential = secret_store.get("provider-key")',"Authenticated workload identity","Scoped provider access","Enterprise buyers pay for governance and reliability.","SECURITY.md"),
    p(68,"PRIVACY & REDACTION","Collect less, retain less, expose less.","Data minimization funnel removing email, phone and secrets",("Classify data before transmission","Redact PII and credentials from logs","Choose region and retention deliberately","Support deletion and access workflows"),'safe_log = redact_mapping(event)',"Request metadata with sensitive fields","Useful telemetry without exposed secrets","Privacy-aware design expands the customers you can serve.","src/ai_api_playbook/safety.py"),
    p(69,"RATE LIMITS & CONCURRENCY","Throughput is a budgeted resource, not an invitation to retry everything at once.","Highway toll gates controlling requests per minute and tokens per minute",("Track request, token and concurrent-job limits","Queue excess work","Apply tenant quotas","Use server hints such as Retry-After"),'semaphore = Limit(concurrency=5)',"Burst of jobs","Smoothed queue with predictable throughput","Capacity control prevents outages and surprise bills.","docs/architecture.md"),
    p(70,"RETRY & BACKOFF","Retry transient failures with patience, jitter and a stopping rule.","Cartoon climber taking wider-spaced steps up a retry hill",("Retry timeouts, connection failures and selected 5xx/429 errors","Do not blindly retry validation or authorization failures","Use exponential backoff with jitter","Protect billable writes with idempotency"),'delay = random(0, min(cap, base*2**attempt))',"Transient failure","Recovered result or controlled final error","Correct retry logic improves reliability without multiplying cost.","src/ai_api_playbook/retry.py"),
    p(71,"CACHING","Do not pay repeatedly for the same stable work.","Refrigerator of reusable approved results with expiration labels",("Cache only when inputs and policy match","Hash normalized prompts and source versions","Set TTL by data freshness","Never leak cached results across tenants"),'key = hash(model + prompt + source_version + policy)',"Repeated deterministic request","Cached validated response","Caching directly improves speed and gross margin.","docs/architecture.md"),
    p(72,"COST CONTROL","Every request needs a maximum acceptable cost before it starts.","Cockpit budget dashboard with per-request, tenant and monthly gauges",("Estimate input before calling","Cap output and agent steps","Track media seconds and tool charges","Stop or downgrade only according to product policy"),'if projected_cost > budget: raise BudgetExceededError',"Planned request + price table","Approved route or budget rejection","Cost-aware architecture makes revenue sustainable.","src/ai_api_playbook/economics.py"),
    p(73,"EVALUATION","Measure the behavior your customer buys.","Lab bench testing outputs against golden cases",("Build representative test cases","Define task-specific rubrics","Compare quality, latency, cost and failure rate","Run evals before model or prompt changes"),'score = evaluate(output, expected, evidence)',"Test dataset + candidate system","Comparable metrics and regression alerts","Evaluation enables confident upgrades and premium guarantees.","src/ai_api_playbook/evaluation.py"),
    p(74,"OBSERVABILITY","Trace one customer outcome across models, tools, retrieval and validation.","Constellation trace linking request spans",("Use correlation IDs across services","Record safe inputs, route, latency, usage and status","Separate provider failure from product failure","Create alerts on success rate and cost per successful job"),'trace_id → model span → tool span → validation span',"Events from the full workflow","Metrics, traces and actionable alerts","Operational visibility reduces downtime and support burden.","docs/architecture.md"),
    p(75,"PROJECT 1 — CITED KNOWLEDGE ASSISTANT","Turn approved documents into answers employees can verify.","Office librarian robot handing cited evidence cards to a staff member",("Buyer: SME operations team","Pain: repeated policy search","Core stack: ingestion, hybrid retrieval, reranking, grounded generation","Success: useful answer with correct source IDs in under 15 seconds"),'answer = rag(question, tenant, permissions)',"Question + authorized corpus","Answer + citations + abstention path","Charge setup + monthly subscription + usage.","projects/blueprints.py"),
    p(76,"KNOWLEDGE ASSISTANT — ECONOMICS","Sell verified time saved, not chatbot messages.","Unit-economics seesaw balancing subscription and variable cost",("Measure answer success rate","Calculate cost per successful answer","Create a human escalation queue","Protect tenant isolation and document permissions"),'margin = price - api - infra - support',"Usage, success and support metrics","Gross margin and improvement plan","Recurring knowledge value creates sticky subscriptions.","src/ai_api_playbook/economics.py"),
    p(77,"PROJECT 2 — DOCUMENT INTELLIGENCE","Convert document queues into validated records and exceptions.","Invoice assembly line with green accepted and amber review lanes",("Buyer: accounting or logistics team","Extract fields with page evidence","Normalize currencies, dates and identifiers","Route low-confidence values to human review"),'document → OCR → schema → rules → review queue',"PDFs and scans","Validated records + exceptions","Price per verified document plus premium review tier.","projects/blueprints.py"),
    p(78,"DOCUMENT INTELLIGENCE — QUALITY LOOP","Accuracy is a workflow of extraction, rules and review.","Feedback flywheel between model, validator and reviewer",("Create field-level ground truth","Track straight-through processing rate","Analyze errors by document type","Never hide uncertainty behind one document-level score"),'field = {value, evidence, confidence, status}',"Extracted candidates","Audited field dataset","Operational data becomes a defensible improvement asset.","projects/publishing_engine/pipeline.py"),
    p(79,"PROJECT 3 — VOICE RECEPTIONIST","Recover missed opportunities without pretending uncertainty is certainty.","Friendly virtual receptionist routing calls to booking, answer or human",("Buyer: clinics and local services","Connect telephony, realtime voice and scheduling tools","Confirm names, dates and contact details","Transfer sensitive or unclear calls"),'call → identify → resolve | book | transfer',"Inbound call","Confirmed outcome + summary","Charge setup, monthly platform fee and included minutes.","projects/blueprints.py"),
    p(80,"VOICE RECEPTIONIST — SAFETY & KPI","Completion matters only when consent, accuracy and escalation work.","Call-center dashboard with completion, transfer, latency and consent gauges",("Track resolved-call rate","Measure booking correctness","Record time to first audio and interruption quality","Audit consent, transfer and failure behavior"),'KPI = successful_verified_outcomes / eligible_calls',"Call events and outcomes","Operational and safety dashboard","Reliability justifies higher-value B2B pricing.","projects/blueprints.py"),
    p(81,"PROJECT 4 — E-COMMERCE CATALOG ENGINE","Transform approved product data into consistent multilingual sales assets.","Product SKU enters a colorful content factory and exits as catalog cards",("Buyer: retailers and exporters","Ground every claim in product data","Generate title, bullets, description and image brief","Validate units, variants, prohibited claims and locale"),'sku_data → schema → localized copy → QA',"Structured SKU + images","Approved channel-ready product pack","Price per SKU, batch or monthly catalog volume.","projects/blueprints.py"),
    p(82,"CATALOG ENGINE — MARGIN DESIGN","Batch stable work; reserve premium models for difficult exceptions.","Tiered conveyor with fast lane, premium lane and human review lane",("Classify SKU complexity first","Cache brand and policy context","Batch routine generations","Measure approval rate and revisions per SKU"),'route = premium if complexity > threshold else efficient',"SKU complexity + policy","Cost-controlled production route","Operational routing protects margins at catalog scale.","src/ai_api_playbook/routing.py"),
    p(83,"PROJECT 5 — LOCALIZATION STUDIO","Create timed, reviewable multilingual media packages.","Podcast waveform crossing four language booths",("Buyer: educators and media teams","Transcribe with speakers and timestamps","Translate for meaning and audience","Generate voice only with rights and consent"),'media → transcript → translation → voice → QA',"Authorized audio/video + target locale","Transcript, subtitle and dubbing package","Charge per source minute plus language and review tiers.","projects/blueprints.py"),
    p(84,"LOCALIZATION — HUMAN REVIEW MAP","Automation accelerates the pipeline; experts protect names, claims and culture.","Language reviewer marking critical zones on a subtitle timeline",("Review proper nouns and terminology","Verify numbers, dates and legal claims","Check timing, pronunciation and cultural fit","Keep source-to-translation traceability"),'segment_id links source → translation → audio',"Timed language segments","Approved localized master","Tiered review becomes a premium quality offering.","projects/blueprints.py"),
    p(85,"PROJECT 6 — PROPERTY MEDIA STUDIO","Generate factual listing packages from approved property evidence.","Property photos entering a studio and exiting as copy, image and video cards",("Buyer: real estate and tourism teams","Extract only observable and supplied facts","Create channel-specific copy and visual briefs","Protect addresses, people, brands and disclosures"),'property_data + approved_media → listing_pack',"Facts, photos and channel","Copy, image variants and short video brief","Sell per-property packages or agency subscriptions.","projects/blueprints.py"),
    p(86,"PROPERTY MEDIA — TRUST CHECKS","A beautiful listing that invents facts is a business liability.","Red factual shield blocking invented amenities from a glossy listing",("Lock factual fields before generation","Flag unsupported adjectives and claims","Review edited images for material changes","Retain approval and source provenance"),'assert every_claim in approved_source',"Generated listing candidate","Factual approved package","Trust and fast turnaround create defensible service value.","SECURITY.md"),
    p(87,"PROJECT 7 — SUPPORT TRIAGE","Classify, retrieve, draft and escalate—without silently sending risky replies.","Support tickets moving into routine, urgent and human lanes",("Buyer: customer support teams","Detect intent, urgency and sentiment cautiously","Retrieve approved policy and account data","Draft with citations; require approval for external send"),'ticket → classify → retrieve → draft → approve/escalate',"Customer message + authorized context","Draft response + route + evidence","Charge per seat plus usage or resolved ticket.","projects/blueprints.py"),
    p(88,"SUPPORT TRIAGE — OPERATING METRICS","Optimize successful resolution, not automated message volume.","Dashboard with resolution, reopen, escalation and cost metrics",("Track first-contact resolution","Measure reopen and correction rate","Audit false urgency and missed escalation","Compare agent time saved with system cost"),'cost_per_resolution = total_variable_cost / verified_resolutions',"Ticket outcomes","Quality and unit-economics dashboard","Outcome metrics prevent automation theater.","src/ai_api_playbook/economics.py"),
    p(89,"PROJECT 8 — PROPOSAL & RFP COPILOT","Make every requirement traceable before drafting persuasive prose.","RFP pages feeding a compliance matrix and proposal storyboard",("Buyer: agencies and B2B vendors","Extract requirements with page citations","Build compliance and evidence matrix","Draft only from approved capabilities and commitments"),'RFP → requirements → evidence matrix → draft → review',"RFP + approved company evidence","Traceable response package","Price per workspace, proposal or annual team license.","projects/blueprints.py"),
    p(90,"PROPOSAL COPILOT — COMMITMENT GATE","The model may draft; an accountable owner must approve every promise.","Contract gatekeeper stamping approved commitments",("Separate facts, plans and commitments","Detect missing evidence","Flag dates, prices and legal language","Record approver and source for final claims"),'commitment.status must equal "approved"',"Draft commitments","Approved proposal or exception list","Governance turns drafting speed into enterprise-grade value.","SECURITY.md"),
    p(91,"PROJECT 9 — NICHE AI TUTOR","Build mastery from an approved curriculum, not unlimited improvisation.","Friendly tutor guiding a learner through concept, practice and feedback stations",("Buyer: training providers","Ground explanations in curated sources","Adapt difficulty from observed performance","Generate quizzes with answer evidence"),'diagnose → teach → practice → assess → adapt',"Learner state + curriculum","Lesson, practice and feedback","Charge per learner/month or license to institutions.","projects/blueprints.py"),
    p(92,"AI TUTOR — LEARNING EVIDENCE","Engagement is not the same as learning.","Progress staircase with pre-test, practice and post-test markers",("Measure pre/post performance","Track misconception categories","Use delayed retrieval practice","Let instructors inspect sources and feedback"),'learning_gain = post_score - pre_score',"Assessment events","Evidence of mastery and next step","Measured learning outcomes support durable educational products.","src/ai_api_playbook/evaluation.py"),
    p(93,"PROJECT 10 — MARKET INTELLIGENCE BRIEF","Convert scattered signals into a cited recurring decision product.","Radar station collecting web signals into a concise executive brief",("Buyer: small strategy teams","Monitor defined sources and topics","Detect what changed since last edition","Separate sourced facts, model inference and analyst judgment"),'collect → deduplicate → compare → synthesize → cite → approve',"Current sources + prior brief","Cited change report","Sell weekly or monthly subscriptions by niche.","projects/blueprints.py"),
    p(94,"MARKET INTELLIGENCE — TRUST & RETENTION","Customers renew when the brief is timely, relevant and auditable.","Subscription flywheel powered by relevance, evidence and actionability",("Track source freshness","Score claim-to-source support","Collect which insights changed decisions","Archive editions and corrections transparently"),'renewal_value = relevance × trust × actionability',"Brief usage and feedback","Improved source and topic policy","A narrow trusted niche beats a generic flood of summaries.","projects/blueprints.py"),
    p(95,"CAPSTONE — AI PUBLISHING ENGINE","Turn an authorized manuscript into a governed publishing package.","A book enters a large multimodal machine and exits as metadata, RAG, audio and launch assets",("User: author or publisher","Input: owned manuscript and metadata","Core: document AI, structured output, RAG and media briefs","Output: human-review package—never automatic publication"),'manuscript → extract → structure → enrich → review → export',"Authorized book/PDF","Publishing and knowledge-commerce package","Offer as author service, publisher tool or SaaS.","projects/publishing_engine/README.md"),
    p(96,"CAPSTONE — DOCUMENT TO KNOWLEDGE","Preserve provenance while turning pages into searchable evidence.","Book pages transforming into traceable chunks connected to a search index",("Fingerprint the source version","Extract headings, tables and page locations","Chunk with section and page metadata","Build retrieval tests from real reader questions"),'source_hash + page + section + chunk_id',"Manuscript pages","Versioned RAG manifest","A reliable book assistant extends the commercial life of content.","projects/publishing_engine/pipeline.py"),
    p(97,"CAPSTONE — MULTILINGUAL COMMERCE PACK","Generate channel-specific assets without losing factual control.","Four-language launch board with Google Books, Zenodo, LinkedIn and catalog cards",("Create schema-valid book metadata","Generate language-specific descriptions and keywords","Prepare audio, image and video briefs","Review facts, rights, locale and platform rules"),'metadata → locale packs → channel templates → approval',"Approved manuscript metadata","Reviewable multilingual launch assets","Sell premium packaging by language and channel.","projects/publishing_engine/pipeline.py"),
    p(98,"CAPSTONE — HUMAN APPROVAL SYSTEM","The final product is not the generation; it is the approved state transition.","Five approval gates protecting a publication button",("Gate 1: factual metadata","Gate 2: rights, likeness and consent","Gate 3: privacy and secret scan","Gate 4: cost and quality","Gate 5: final publication authority"),'if all(gates): export(); else: review_queue()',"Generated package + evidence","Approved export or explicit exceptions","Human approval preserves trust while automation creates speed.","projects/publishing_engine/pipeline.py"),
    p(99,"RELEASE, CITATION & DOI","Freeze a verified version so others can run, inspect and cite the same work.","Git tag becoming a Zenodo archive and DOI citation card",("Test and document before tagging","Use semantic versioning and a clear changelog","Include CITATION.cff, license and ORCID","Create separate linked records for software and book"),'git tag v1.0.0 → GitHub Release → Zenodo DOI',"Verified source tree","Citable archived software release","A DOI strengthens discoverability and scholarly attribution.","docs/zenodo-release.md"),
    p(100,"FROM API CALL TO DURABLE VALUE","The future belongs to systems that combine intelligence with evidence, control and economics.","Cinematic staircase from first API call to trusted product and recurring value",("Begin with one measurable customer pain","Build a safe, observable workflow","Evaluate quality and cost per successful job","Publish evidence, improve continuously and earn trust"),'capability + workflow + evidence + governance + distribution = value',"A useful idea and disciplined engineering","A trusted product—not merely a demo","The API is rented intelligence; the system around it is your lasting asset.","README.md"),
]

assert len(PAGES) == 100
assert [page.number for page in PAGES] == list(range(1, 101))

PALETTES = [
    ("electric blue, cyan, white and deep navy", "#1368CE"),
    ("coral, warm yellow, cream and charcoal", "#F05D5E"),
    ("violet, mint, white and midnight purple", "#6C4CE6"),
    ("teal, lime accents, pale gray and black", "#008C8C"),
    ("orange, sky blue, white and graphite", "#F28C28"),
]

LAYOUTS = [
    "radial capability map",
    "modular bento-grid",
    "left-to-right process pipeline",
    "top-down decision tree",
    "comic-strip sequence",
    "dashboard with gauges and cards",
    "isometric system cutaway",
    "split-screen comparison",
    "subway-map information flow",
    "central metaphor with orbiting notes",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(8.5)
    normal.font.color.rgb = RGBColor(35, 43, 58)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    for name, size, color, before, after in (
        ("Title", 25, RGBColor(13, 45, 82), 0, 8),
        ("Subtitle", 12, RGBColor(80, 93, 115), 0, 10),
        ("Heading 1", 15, RGBColor(19, 104, 206), 10, 5),
        ("Heading 2", 10.5, RGBColor(13, 45, 82), 6, 3),
    ):
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Heading 2" else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "AI APIs IN PRACTICE  |  VISUAL PROMPT MANUSCRIPT"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(7)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor(99, 115, 139)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Faramarz Kowsari  •  First Edition 2026")
    footer.runs[0].font.size = Pt(7)
    footer.runs[0].font.color.rgb = RGBColor(99, 115, 139)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI APIs IN PRACTICE")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(13, 45, 82)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Visual Infographic Production Manuscript\nFront Cover + 100 Numbered Pages + Back Cover")
    sr.font.name = "Aptos"
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(19, 104, 206)
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = author.add_run("Faramarz Kowsari")
    ar.bold = True
    ar.font.size = Pt(16)
    ar.font.color.rgb = RGBColor(240, 93, 94)
    doc.add_page_break()


def exact_copy(page: PageSpec) -> str:
    bullets = "\n".join(f"• {point}" for point in page.points)
    return (
        f'TITLE: “{page.title}”\n'
        f'SUBTITLE: “{page.subtitle}”\n'
        f'GOLDEN NOTES:\n{bullets}\n'
        f'CODE CARD:\n{page.code}\n'
        f'INPUT CARD: “{page.input_text}”\n'
        f'OUTPUT CARD: “{page.output_text}”\n'
        f'COMMERCIAL LENS: “{page.commercial}”\n'
        f'REPOSITORY: “{page.repo_path}”\n'
        f'FOOTER BRAND: “AI APIs in Practice • Faramarz Kowsari”\n'
        f'PAGE NUMBER: “{page.number:03d}”'
    )


def build_prompt(page: PageSpec) -> str:
    palette, accent = PALETTES[(page.number - 1) % len(PALETTES)]
    layout = LAYOUTS[(page.number - 1) % len(LAYOUTS)]
    return f"""Create one independent, print-ready A4 portrait educational infographic, 2480 × 3508 px, 300 DPI, for page {page.number:03d} of the English book “AI APIs in Practice” by Faramarz Kowsari. This is a technical publishing page, not a poster collage and not a screenshot. Use a bright editorial background, generous safe margins, crisp vector-like shapes, subtle paper texture, precise alignment and strong information hierarchy.

UNIQUE VISUAL DIRECTION
Build the page as a {layout}. The central conceptual metaphor is: {page.visual}. Use the palette {palette}, with {accent} as the dominant accent. Add small original cartoon engineering characters, clean icons, arrows, data packets, code windows, evidence tags and hand-drawn annotation marks only where they clarify meaning. Make the visual metaphor do real explanatory work. Do not use generic glowing brains, random circuit backgrounds or decorative clutter.

TYPOGRAPHY
Use a maximum of three complementary type families: a bold geometric sans-serif for the title, a highly legible humanist sans-serif for explanatory text, and a monospaced developer font for code. Add a restrained handwritten marker style for only two or three short annotations such as “golden rule”, “verify”, or “cost check”. Vary size, weight and color deliberately. Keep every character readable at A4 print size. Never create tiny paragraphs. Render all supplied wording exactly, with correct spelling, punctuation and capitalization. Do not invent extra text.

INFORMATION ARCHITECTURE
Place the title at the top, subtitle directly beneath it, the explanatory visual in the center, four golden-note cards around or beside the visual, and a lower technical strip containing CODE, INPUT, OUTPUT, COMMERCIAL LENS and REPOSITORY. The code must appear as real monospaced code with preserved line breaks, not as decorative pseudo-text. Use arrows to make input → processing → output relationships instantly understandable. Reserve the bottom safe zone for the brand line and a clearly visible three-digit page number.

EXACT VISIBLE COPY
{exact_copy(page)}

QUALITY AND SAFETY RULES
No logos unless supplied as official reference assets. Provider names may appear only when present in the exact copy. No QR code, barcode, ISBN, watermark, fake URL, illegible microtext, random letters, duplicated cards, cropped text, incorrect code, distorted hands, or meaningless UI. Keep all text inside safe margins. Ensure high contrast, balanced whitespace and a clean professional finish appropriate for an internationally published technical book. Output exactly one standalone infographic image for page {page.number:03d}."""


FRONT_COVER_PROMPT = """Create one independent print-ready A4 portrait front cover, 2480 × 3508 px, 300 DPI, for an international technical book. Bright sophisticated editorial design: a human software engineer stands at the center of a precise visual ecosystem where text, code, tools, vector search, documents, image, voice and video flow through connected API gateways into real products. Mix clean vector infographic language with subtle cinematic depth; white and pale-blue base, electric blue, violet, coral and mint accents. Add small hand-drawn arrows and annotation sparks without clutter. Title at the top in large geometric sans-serif: “AI APIs IN PRACTICE”. Subtitle: “A Visual Guide to Multimodal, Agentic, and Revenue-Ready Applications”. Place “Faramarz Kowsari” prominently at the bottom. Use a monospaced micro-accent reading “TEXT • TOOLS • RAG • VISION • VOICE • VIDEO”. No QR code, barcode, ISBN, DOI, fake logos or extra text. Exact spelling, generous margins, premium global technology-publishing quality."""

BACK_COVER_PROMPT = """Create one independent print-ready A4 portrait back cover, 2480 × 3508 px, 300 DPI, visually continuous with the front cover: bright pale background, connected API streams and small technical cartoon elements fading into generous readable space. Include a tasteful author portrait area using the supplied reference portrait when provided; preserve recognizable identity and natural professional appearance. Large heading: “FROM API CALLS TO TRUSTED PRODUCTS”. Promotional copy: “AI APIs in Practice is a visual engineering guide to building with language models, tools, agents, RAG, documents, images, speech, realtime voice and video. Across 100 information-rich infographics, the book connects every capability to production architecture, security, evaluation, cost control and commercially realistic product design. Each page maps to runnable companion code, helping readers move from a first request to systems that can be tested, measured, cited and improved.” Add a small feature list: “100 Visual Lessons • Runnable Companion Code • 10 Commercial Blueprints • One Integrated Publishing Capstone”. Add author line: “Faramarz Kowsari — Author, Software Engineer and AI Researcher”. No QR code, barcode, ISBN, DOI, price, fake testimonial or extra text. Use large readable typography, exact spelling and premium international publishing quality."""


def add_prompt_page(doc: Document, label: str, prompt: str, *, accent: str = "1368CE") -> None:
    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run(label)
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(17.7)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F7FC")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(prompt)
    run.font.name = "Aptos"
    run.font.size = Pt(8.2)
    run.font.color.rgb = RGBColor(31, 45, 61)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(0)
    nr = note.add_run("Production check: one image only • exact visible copy • A4 portrait • no invented text")
    nr.bold = True
    nr.font.size = Pt(7.5)
    nr.font.color.rgb = RGBColor.from_string(accent)


def build(output_path: Path) -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    doc.add_heading("Production System", level=1)
    doc.add_paragraph(
        "This manuscript contains self-contained English prompts for the front cover, 100 numbered infographic pages, and the back cover. Each numbered page maps to the companion GitHub repository. Generate pages as independent images in batches of ten; preserve filenames as page-001.png through page-100.png."
    )
    doc.add_heading("Global Acceptance Checklist", level=2)
    for item in (
        "Exactly one A4 portrait infographic per prompt; no collage or multi-page image.",
        "All visible text and code match the supplied copy exactly.",
        "Page number is present at the bottom in three-digit format.",
        "Readable at normal print size; no microtext, clipping or overlap.",
        "Visual metaphor clarifies the concept and layout varies across pages.",
        "No QR code, barcode, ISBN, fake DOI, watermark or fabricated provider logo.",
        "Repository path and author brand remain consistent.",
    ):
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()

    add_prompt_page(doc, "FRONT COVER PROMPT", FRONT_COVER_PROMPT, accent="F05D5E")
    doc.add_page_break()
    for page in PAGES:
        add_prompt_page(doc, f"PAGE {page.number:03d} — {page.title}", build_prompt(page), accent=PALETTES[(page.number - 1) % len(PALETTES)][1].lstrip("#"))
        if page.number != 100:
            doc.add_page_break()
    doc.add_page_break()
    add_prompt_page(doc, "BACK COVER PROMPT", BACK_COVER_PROMPT, accent="6C4CE6")
    doc.core_properties.title = "AI APIs in Practice — Visual Infographic Prompt Manuscript"
    doc.core_properties.subject = "Front cover, 100 infographic prompts and back cover"
    doc.core_properties.author = "Faramarz Kowsari"
    doc.core_properties.keywords = "AI APIs, infographic book, prompt manuscript, RAG, agents, multimodal AI"
    doc.core_properties.comments = "First Edition 2026"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    build(Path("AI_APIs_in_Practice_Infographic_Prompt_Manuscript.docx"))
